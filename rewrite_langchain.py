from langchain.messages import HumanMessage, SystemMessage
from langchain_chroma import Chroma
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
import streamlit as st
from langchain_core.output_parsers import StrOutputParser

list_paragraphs=[]


if "messages" not in st.session_state:
    st.session_state.messages = []
@st.cache_resource    
def loadresume(resume):
 try:
    embeddings = HuggingFaceEmbeddings(model="sentence-transformers/all-MiniLM-L6-v2")
    doc = DocxDocument(resume)
    print("Resume loaded successfully")
    for paragraph in doc.paragraphs:
     list_paragraphs.append(Document(page_content=paragraph.text))
    vector_store =Chroma.from_documents(
                   documents=list_paragraphs,
                   embedding=embeddings)
    
    return vector_store
    
 except OSError as e:
     print(f"OSError occurred while loading the resume, {e}")
     return None
 except Exception as e:
     print(f"Exception occurred while loading the resume, {e}")
     return None

def InvokeOllama(prompt,vector_store):
    embed_query=vector_store.similarity_search(prompt,k=2)
    context = "\n\n".join([doc.page_content for doc in embed_query])  
    
    
    template = ChatPromptTemplate.from_messages([
          ("system", "You are a resume assistant. Answer only based on this resume context {context}"),
          ("human", "{question}")
       ])
    
    chain = template | model | parser
    response = chain.invoke({"context":context, "question":prompt})
    return response

parser = StrOutputParser()
model = ChatOllama(model="llama3.2")
vector=loadresume("/Users/hravs/Python_projects/ResumeBot/data/Test_Resume_Sample.docx")
# message= HumanMessage(get_completion("what is the name of candidate?",vector))
# response =model.invoke([message])
# print(response.content)


for history in st.session_state.messages:
    if history["role"] == "system":continue
    if history["role"]:
        with st.chat_message(history["role"]):
             st.write(history["content"])
user = st.chat_input("hello , your chatbot is here")
if user:
    result=InvokeOllama(user,vector)

    st.session_state.messages.append({"role":"user","content":user})
    st.session_state.messages.append({"role":"assistant","content":result})
    st.rerun()

# notes:
# DocxDocument reads raw text → LangChain Document wraps it which chroma understands → ChromaDB stores it with embeddings → Similarity Search retrieves relevant chunks → SystemMessage passes context to LLM → Ollama generates the answer → Streamlit displays it

# RAG = Retrieval Augmented Generation

# Retrieval = similarity search in ChromaDB
# Augmented = adding that context to the prompt
# Generation = LLM generating the answer