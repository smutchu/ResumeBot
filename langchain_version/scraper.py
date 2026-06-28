import requests
from bs4 import BeautifulSoup
from langchain_chroma import Chroma
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
import streamlit as st
from langchain_core.output_parsers import StrOutputParser
list_paragraphs=[]
from langchain.tools import tool
def loadresume(resume):
 try:
    embeddings = HuggingFaceEmbeddings(model="sentence-transformers/all-MiniLM-L6-v2")
    doc = DocxDocument(resume)
    print("Resume loaded successfully")
    for paragraph in doc.paragraphs:
     list_paragraphs.append(Document(page_content=paragraph.text))
    vector_store =Chroma.from_documents(
                   documents=list_paragraphs,
                   embedding=embeddings)
    
    return vector_store
    
 except OSError as e:
     print(f"OSError occurred while loading the resume, {e}")
     return None
 except Exception as e:
     print(f"Exception occurred while loading the resume, {e}")
     return None
 
vector=loadresume("/Users/hravs/Python_projects/ResumeBot/data/Test_Resume_Sample.docx")
model = ChatOllama(model="llama3.2")
@tool(description="Scrapes a job description from the given URL. Falls back to manual paste if scraping fails.")
def scrape_job_description(url: str) -> str:
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()

        # Parse the HTML
        soup = BeautifulSoup(response.text, "html.parser")

        jd = soup.find_all("div", class_="article__content")[2].get_text(strip=True)

        return jd

    except Exception as e:
        print(f"Error: {e}")
        
        return input("please paste the JD, couldn't parse the url: ")


 
parser = StrOutputParser()

@tool(description="Scores the resume against the job description out of 10 with reasoning.")
def get_score(jd: str) -> str:
   embed_query=vector.similarity_search(jd,k=10)
   context = "\n\n".join([doc.page_content for doc in embed_query])  
   template = ChatPromptTemplate.from_messages([
          ("system", "You are a resume fit scorer. Compare this resume {context} to this job description {jd}. Your FIRST line must be exactly 'Score: X/10' where X is a single number. Nothing else on that line. Then on the next line explain what matches and what is missing."),
          ("human", "Please score my fit for this job.")
       ])
   chain = template | model | parser
   response = chain.invoke({"context":context, "jd":jd})
   return response

@tool(description="Rewrites the resume to match the job description and be ATS friendly.")
def modify_resume(jd: str) -> str:
   embed_query=vector.similarity_search(jd,k=10)
   context = "\n\n".join([doc.page_content for doc in embed_query])  
   template = ChatPromptTemplate.from_messages([
          ("system", "You are a resume editor. Rewrite the ENTIRE resume below word for word, replacing and tailoring the content to match the job description. Do NOT give advice or suggestions. Do NOT explain what to change. Just output the fully rewritten resume text and nothing else. Here is the resume: {context}. Here is the job description: {jd}."),
          ("human", "Rewrite my resume.")
       ])
   chain = template | model | parser
   response = chain.invoke({"context":context, "jd":jd})
   return response
# JD = scrape_job_description.invoke({"url": input("please enter url: ")})
   

# score= get_score.invoke({"jd":JD})
# print(score.split()[1].split('/')[0])

# score_number = modify_resume.invoke({"jd":JD,}) if float(score.split()[1].split('/')[0]) >= 7 else "not a good fit"

# print(score_number)

@tool(description="Saves the modified resume as a docx file.")
def save_document(response_from_llm: str) -> str:
   doc = DocxDocument()
   for line in response_from_llm.split("\n"):
      doc.add_paragraph(line)
   doc.save("modified.docx")
   
# save_document.invoke({"response_from_llm":score_number})  
      
# agent = create_agent(
#     model=model,
#     tools=[scrape_job_description, get_score, modify_resume, save_document],
#     system_prompt="You are a resume assistant. Help the user score and tailor their resume to job descriptions."
# )
# response = agent.invoke({
#     "messages": [{"role": "user", "content": "Score my resume against this JD: <url>"}]
# })
   
   


